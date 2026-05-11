import json
import os
import sys
from typing import Any

import docx
import pdfplumber
from google import genai
from google.genai import types
from ollama import chat

PROMPT_TEMPLATE = """
You are a strict JSON generator. Your task is to extract information from a CV.

RULES (MUST FOLLOW):
- Output ONLY valid JSON. No explanation, no markdown, no extra text.
- Do NOT add any fields that are not in the schema.
- Do NOT rename fields.
- Do NOT infer or guess missing information.
- If a field is missing:
  - "" for strings
  - [] for arrays
- Keep extracted text concise and factual.
- Do NOT include labels like "Name:", "Email:", etc.
- Do NOT hallucinate companies, dates, or skills.
- Use double quotes for JSON keys and string values.

SCHEMA:
{{
  "name": "",
  "email": "",
  "phone": [],
  "location": "",
  "summary": "",
  "skills": [],
  "experience": [
    {{
      "type": "job" | "internship" | "freelance" | "project",
      "title": "",
      "company": "",
      "start_date": "",
      "end_date": "",
      "description": "",
      "tech": []
    }}
  ],
  "education": [
    {{
      "degree": "",
      "institution": "",
      "year": ""
    }}
  ],
  "languages": [],
  "certifications": [],
  "references": []
}}

EXTRACTION INSTRUCTIONS:

GENERAL:
- Extract exactly what is written in the CV.
- Do NOT assume missing data.
- If unsure, leave empty.

EXPERIENCE (IMPORTANT):
- Include ALL work-related items under "experience":
  - jobs
  - internships
  - freelance work
  - projects (VERY IMPORTANT)
- Each entry must be a separate object.
- Set "type" field to one of:
  - "job"
  - "internship"
  - "freelance"
  - "project"
- If unclear whether something is a job or project, classify it as "project".
- One object per distinct item (not per bullet point).
- If no experience exists, return [].

PROJECT RULES (inside experience):
- Projects MUST be included under experience with type = "project".
- Projects may come from sections like:
  - Projects
  - Personal Projects
  - Portfolio
  - Hackathons
  - GitHub repositories
- Projects usually do NOT have a company name; keep company as "" if missing.

DESCRIPTION RULES:
- Each bullet must be preserved separately.
- Do NOT merge bullets into one string.

TECH RULES:
- Extract only explicitly mentioned technologies.
- Do NOT infer skills.

EDUCATION:
- One object per degree.

SKILLS:
- Extract only explicitly written skills.
- No inference.

CV TEXT:
{text}
"""


def extract_text(filepath) -> str | None:
    """Extract text from PDF or DOCX files."""
    if filepath.endswith(".pdf"):
        with pdfplumber.open(filepath) as pdf:
            return "\n".join(page.extract_text() for page in pdf.pages)
    elif filepath.endswith(".docx"):
        doc = docx.Document(filepath)
        text = []

        # Extract text from paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text.append(p.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for r in table.rows:
                row = []
                for cell in r.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row.append(cell_text)
                if row:
                    text.append(" | ".join(row))

        return "\n".join(text)
    else:
        return None


def parse_cv_ollama(text) -> dict[str, Any]:
    """Parse CV text into structured JSON according to the defined schema using Ollama."""
    prompt = PROMPT_TEMPLATE.format(text=text)

    response = chat(
        model="gemma3",
        messages=[
            {
                "role": "system",
                "content": "You output ONLY JSON that strictly follows the schema. No variation allowed.",
            },
            {"role": "user", "content": prompt},
        ],
        format="json",
        think=False,
        options={
            "num_ctx": 4096,
            "temperature": 0,
            "top_p": 1,
            "top_k": 1,
        },
    )

    if not response.done:
        raise ValueError("Model did not return a response.")

    raw = response.message.content
    if raw is None:
        raise ValueError("Model response is empty.")

    return json.loads(raw.strip())


def parse_cv_genai(text) -> dict[str, Any]:
    """Parse CV text using GenAI's structured output capabilities."""
    prompt = PROMPT_TEMPLATE.format(text=text)
    client = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))
    model = "gemini-2.5-flash"
    response = client.models.generate_content(
        model=model,
        contents=prompt,
        config=types.GenerateContentConfig(response_mime_type="application/json", temperature=0),
    )

    raw = response.text
    if raw is None:
        raise ValueError("Model response is empty.")
    return json.loads(raw.strip())


def parse_cv(text) -> dict[str, Any]:
    """Parse CV text into structured JSON. Uses GenAI if GOOGLE_API_KEY is set, otherwise Ollama."""
    if os.getenv("GOOGLE_API_KEY"):
        return parse_cv_genai(text)
    return parse_cv_ollama(text)


def main():
    if len(sys.argv) != 2:
        print("Usage: python cv_parse.py <path_to_cv>")
        return

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"File {filepath} does not exist.")
        return

    text = extract_text(filepath)
    data = json.dumps(parse_cv(text))
    filename = os.path.splitext(os.path.basename(filepath))[0]
    output_path = f"{filename}_parsed.json"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(data)

    print(f"Output file: {output_path}")


if __name__ == "__main__":
    main()
